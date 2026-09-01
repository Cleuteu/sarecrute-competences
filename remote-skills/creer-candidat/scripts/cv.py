#!/usr/bin/env python3
"""Extrait le texte d'un CV, et le joint à une fiche candidat en pièce jointe.

Trois gestes indépendants :

    python3 cv.py --find "Réhane Chiron Gonnon"
        → cherche le fichier du CV sur le disque (Téléchargements, Bureau, Documents…)
          quand le recruteur a déposé le CV dans la conversation sans donner de chemin.

    python3 cv.py --extract "/chemin/CV Dupont.pdf"
        → écrit le texte extrait sur stdout, à recopier tel quel dans « CV text ».

    python3 cv.py --upload "/chemin/CV Dupont.pdf" --record recXXXXXXXXXXXXXX
        → téléverse le fichier dans le champ pièce jointe « CV » du candidat.

Le téléversement passe par l'API REST (content.airtable.com) : le MCP Airtable ne sait
pas écrire de pièce jointe. Il exige AIRTABLE_API_KEY dans l'environnement.

L'extraction essaie pdftotext (poppler), puis pypdf, puis python-docx pour un .docx.
Si aucun n'est disponible, elle sort en code 3 avec la consigne de repli : lire le
fichier avec l'outil Read et le retranscrire fidèlement.
"""

import argparse
import base64
import json
import mimetypes
import os
import shutil
import subprocess
import sys
import time
import urllib.error
import urllib.request

BASE_ID = "appP0W2ISytaNyAhG"
CHAMP_CV = "fldsj2ukrpjPzu5ox"  # Candidats › CV (pièces jointes)
TAILLE_MAX = 5 * 1024 * 1024  # limite de l'endpoint uploadAttachment


def erreur(message, code=2, **extra):
    print(json.dumps({"erreur": message, **extra}, ensure_ascii=False, indent=2))
    sys.exit(code)


# ---------------------------------------------------------------------------- recherche

DOSSIERS = ["~/Downloads", "~/Desktop", "~/Documents", "."]
EXTENSIONS = (".pdf", ".docx", ".doc", ".odt", ".rtf", ".pages", ".txt")
AGE_MAX_JOURS = 60


def sans_accents(s):
    import unicodedata

    s = unicodedata.normalize("NFKD", s or "")
    return "".join(c for c in s if not unicodedata.combining(c)).lower()


def cmd_find(nom):
    """Un CV déposé dans la conversation n'a pas toujours de chemin exploitable ;
    il est presque toujours quelque part sur le disque, sous un nom qui contient
    celui du candidat. On propose, on ne choisit pas."""
    jetons = [t for t in sans_accents(nom).replace("-", " ").split() if len(t) > 2]
    maintenant = time.time()
    trouves = []

    for d in DOSSIERS:
        d = os.path.abspath(os.path.expanduser(d))
        if not os.path.isdir(d):
            continue
        try:
            entrees = os.listdir(d)
        except PermissionError:
            continue
        for nom_fichier in entrees:
            chemin = os.path.join(d, nom_fichier)
            if not os.path.isfile(chemin):
                continue
            if not nom_fichier.lower().endswith(EXTENSIONS):
                continue
            age_j = (maintenant - os.path.getmtime(chemin)) / 86400
            if age_j > AGE_MAX_JOURS:
                continue
            plat = sans_accents(nom_fichier)
            touches = [t for t in jetons if t in plat]
            if not touches and "cv" not in plat.split(".")[0].replace("_", " ").split():
                continue
            trouves.append(
                {
                    "chemin": chemin,
                    "jetons_du_nom_trouves": touches,
                    "modifie_il_y_a_jours": round(age_j, 1),
                    "octets": os.path.getsize(chemin),
                }
            )

    # Le nom du candidat prime sur la fraîcheur : un « CV.pdf » récent d'un autre
    # candidat ne doit pas passer devant « CV Chiron Gonnon.pdf » d'il y a trois jours.
    trouves.sort(key=lambda f: (-len(f["jetons_du_nom_trouves"]), f["modifie_il_y_a_jours"]))

    print(
        json.dumps(
            {
                "recherche": nom,
                "jetons": jetons,
                "dossiers_explores": [os.path.abspath(os.path.expanduser(d)) for d in DOSSIERS],
                "nb_resultats": len(trouves),
                "resultats": trouves[:10],
                "consigne": "ne jamais téléverser un fichier dont aucun jeton du nom ne "
                "correspond sans le faire confirmer par le recruteur — joindre le CV de "
                "quelqu'un d'autre à une fiche est une fuite de données",
            },
            ensure_ascii=False,
            indent=2,
        )
    )


# --------------------------------------------------------------------------- extraction


def extraire_pdf(chemin):
    if shutil.which("pdftotext"):
        # -layout préserve les colonnes : un CV en deux colonnes devient illisible sans.
        r = subprocess.run(
            ["pdftotext", "-layout", "-enc", "UTF-8", chemin, "-"],
            capture_output=True,
            text=True,
        )
        if r.returncode == 0 and r.stdout.strip():
            return r.stdout, "pdftotext -layout"
    try:
        from pypdf import PdfReader
    except ImportError:
        return None, None
    pages = [(p.extract_text() or "") for p in PdfReader(chemin).pages]
    texte = "\n".join(pages)
    return (texte, "pypdf") if texte.strip() else (None, None)


def extraire_docx(chemin):
    try:
        import docx
    except ImportError:
        return None, None
    d = docx.Document(chemin)
    morceaux = [p.text for p in d.paragraphs]
    for t in d.tables:  # les CV mettent souvent les dates dans un tableau
        for ligne in t.rows:
            morceaux.append("\t".join(c.text for c in ligne.cells))
    texte = "\n".join(morceaux)
    return (texte, "python-docx") if texte.strip() else (None, None)


def extraire(chemin):
    ext = os.path.splitext(chemin)[1].lower()
    if ext in (".txt", ".md", ".text"):
        with open(chemin, encoding="utf-8", errors="replace") as f:
            return f.read(), "lecture directe"
    if ext == ".pdf":
        return extraire_pdf(chemin)
    if ext in (".docx", ".dotx"):
        return extraire_docx(chemin)
    return None, None


def cmd_extract(chemin):
    if not os.path.isfile(chemin):
        erreur(f"fichier introuvable : {chemin}")
    texte, outil = extraire(chemin)
    if not texte:
        erreur(
            "aucun extracteur disponible pour ce fichier, ou fichier sans couche texte "
            "(CV scanné en image)",
            code=3,
            fichier=chemin,
            repli="lire le fichier avec l'outil Read, puis retranscrire fidèlement dans "
            "« CV text » — sans reformuler, sans réordonner, sans corriger",
            pour_installer_un_extracteur="brew install poppler  (pdftotext)",
        )
    print(f"# {len(texte)} caractères extraits par {outil} — à recopier tel quel", file=sys.stderr)
    print(texte)


# ------------------------------------------------------------------------------ upload


def cmd_upload(chemin, record_id):
    if not os.path.isfile(chemin):
        erreur(f"fichier introuvable : {chemin}")
    if not record_id.startswith("rec"):
        erreur(f"recordId invalide : {record_id!r} — attendu recXXXXXXXXXXXXXX")

    taille = os.path.getsize(chemin)
    if taille > TAILLE_MAX:
        erreur(
            f"fichier trop gros ({taille / 1e6:.1f} Mo) — l'endpoint plafonne à 5 Mo",
            fichier=chemin,
            repli="joindre le fichier à la main depuis l'interface Airtable",
        )

    cle = os.environ.get("AIRTABLE_API_KEY")
    if not cle:
        erreur(
            "AIRTABLE_API_KEY absente de l'environnement",
            repli="joindre le CV à la main depuis l'interface Airtable — « CV text » suffit "
            "à l'enrichissement, la pièce jointe est pour le confort du recruteur",
        )

    with open(chemin, "rb") as f:
        contenu = f.read()
    type_mime = mimetypes.guess_type(chemin)[0] or "application/octet-stream"

    url = (
        f"https://content.airtable.com/v0/{BASE_ID}/{record_id}/{CHAMP_CV}/uploadAttachment"
    )
    corps = json.dumps(
        {
            "contentType": type_mime,
            "file": base64.b64encode(contenu).decode("ascii"),
            "filename": os.path.basename(chemin),
        }
    ).encode("utf-8")

    req = urllib.request.Request(
        url,
        data=corps,
        headers={"Authorization": f"Bearer {cle}", "Content-Type": "application/json"},
        method="POST",
    )
    try:
        with urllib.request.urlopen(req, timeout=120) as r:
            reponse = json.loads(r.read().decode("utf-8"))
    except urllib.error.HTTPError as e:
        detail = e.read().decode("utf-8", "replace")[:500]
        erreur(
            f"Airtable a refusé le téléversement (HTTP {e.code})",
            detail=detail,
            repli="joindre le CV à la main depuis l'interface — ne pas réessayer en boucle",
        )
    except Exception as e:
        erreur(f"téléversement impossible : {e}")

    pieces = reponse.get("fields", {}).get(CHAMP_CV, [])
    print(
        json.dumps(
            {
                "ok": True,
                "record": record_id,
                "fichier": os.path.basename(chemin),
                "octets": taille,
                "pieces_jointes_sur_la_fiche": len(pieces),
            },
            ensure_ascii=False,
            indent=2,
        )
    )


def main():
    p = argparse.ArgumentParser(description=__doc__)
    p.add_argument("--find", metavar="NOM", help="chercher le fichier du CV sur le disque")
    p.add_argument("--extract", metavar="FICHIER", help="extraire le texte sur stdout")
    p.add_argument("--upload", metavar="FICHIER", help="téléverser en pièce jointe")
    p.add_argument("--record", metavar="recXXX", help="recordId du candidat (avec --upload)")
    args = p.parse_args()

    if args.find:
        cmd_find(args.find)
    elif args.extract:
        cmd_extract(args.extract)
    elif args.upload:
        if not args.record:
            erreur("--upload exige --record recXXXXXXXXXXXXXX")
        cmd_upload(args.upload, args.record)
    else:
        p.print_help()
        sys.exit(1)


if __name__ == "__main__":
    main()
